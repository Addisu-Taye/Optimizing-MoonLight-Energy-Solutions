import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# 1. Load and combine data
benin = pd.read_csv("data/raw/benin-malanville.csv", parse_dates=["Timestamp"])
togo = pd.read_csv("data/raw/togo-dapaong_qc.csv", parse_dates=["Timestamp"])
sierra = pd.read_csv("data/raw/sierraleone-bumbuna.csv", parse_dates=["Timestamp"])

benin["Country"] = "Benin"
togo["Country"] = "Togo"
sierra["Country"] = "Sierra Leone"
df = pd.concat([benin, togo, sierra], ignore_index=True)

# 2. Basic clean
df = df.fillna(df.median(numeric_only=True))

# 3. Create folders
os.makedirs("unified_figs", exist_ok=True)

# 4. Create Word doc
doc = Document()
doc.add_heading("🌍 Unified Solar Profiling Report: Benin, Togo & Sierra Leone", 0)
doc.add_paragraph("This document summarizes and compares solar datasets across Benin, Togo, and Sierra Leone.")

# Summary statistics
doc.add_heading("Summary Statistics", level=1)
summary = df.groupby("Country")[["GHI", "DNI", "DHI"]].agg(["mean", "median", "std", "min", "max"])
doc.add_paragraph(summary.round(2).to_string())

# Correlation heatmap
sns.heatmap(df[["GHI", "DNI", "DHI", "TModA", "TModB"]].corr(), annot=True, cmap="coolwarm")
plt.title("Correlation Heatmap")
plt.tight_layout()
plt.savefig("unified_figs/corr.png")
plt.close()
doc.add_picture("unified_figs/corr.png", width=Inches(5.5))
doc.add_paragraph("Figure: Correlation matrix across combined dataset.")

# Monthly GHI trend
df["Month"] = df["Timestamp"].dt.month
monthly = df.groupby(["Country", "Month"])["GHI"].mean().reset_index()
sns.lineplot(data=monthly, x="Month", y="GHI", hue="Country", marker="o")
plt.title("Monthly GHI Trends")
plt.tight_layout()
plt.savefig("unified_figs/monthly_ghi.png")
plt.close()
doc.add_picture("unified_figs/monthly_ghi.png", width=Inches(5.5))

# GHI Boxplot
sns.boxplot(data=df, x="Country", y="GHI")
plt.title("GHI Distribution by Country")
plt.tight_layout()
plt.savefig("unified_figs/ghi_boxplot.png")
plt.close()
doc.add_picture("unified_figs/ghi_boxplot.png", width=Inches(5.5))

# Save document
doc.save("Unified_Solar_Profiling_Report.docx")
print(" Report saved as: Unified_Solar_Profiling_Report.docx")
